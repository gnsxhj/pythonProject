#!/usr/bin/env python
# -*- coding:utf-8 -*-
# coding=utf-8

# ***********************************************************************************************************
#import docx
#def info_update(doc,old_info, new_info):
#    '''此函数用于批量替换合同中需要替换的信息
#    doc:文件
#    old_info和new_info：原文字和需要替换的新文字
#    '''
#    #读取段落中的所有run，找到需替换的信息进行替换
#    for para in doc.paragraphs: #
#        for run in para.runs:
#            print(run.text)
#            run.text = run.text.replace(old_info, new_info) #替换信息
#    #读取表格中的所有单元格，找到需替换的信息进行替换
#    for table in doc.tables:
#        for row in table.rows:
#            for cell in row.cells:
#                cell.text = cell.text.replace(old_info, new_info) #替换信息
#
#import os #用于获取目标文件所在路径
#path="c:/Users/XUHAIJUN/Documents/SVW_Consulter_Strategie_2023年/人员简历/test/" # 文件夹路径
#new_path = "c:/Users/XUHAIJUN/Documents/SVW_Consulter_Strategie_2023年/人员简历/test_new/"
#files=[]
#for file in os.listdir(path):
#    if file.endswith(".docx"): #排除文件夹内的其它干扰文件，只获取word文件
#        files.append(path+file)
#
#for file in files:
#    doc = docx.Document(file)
#    info_update(doc,"信息技术", "软件科技")
#    doc.save(new_path+"{}".format(file.split("/")[-1]))
#    print("{}替换完成".format(file))
# ***********************************************************************************************************
import os
import docx
from docx import Document

# 放了一些docx 文件
old_file_path = "c:/Users/XUHAIJUN/SynologyDrive/DTAS/项目信息/一汽大众/2024年框架/人员简历/"
# 生成新文件后的存放地址
new_file_path = "c:/Users/XUHAIJUN/SynologyDrive/DTAS/项目信息/一汽大众/2024年框架/new/"

replace_dict = {
    "君昱": "棣拓",
}

def check_and_change(document, replace_dict):
    """
    遍历word中的所有 paragraphs，在每一段中发现含有key 的内容，就替换为 value 。
    （key 和 value 都是replace_dict中的键值对。）
    """
    for para in document.paragraphs:
        for i in range(len(para.runs)):
            for key, value in replace_dict.items():
                if key in para.runs[i].text:
                    print(key+"-->"+value)
                    para.runs[i].text = para.runs[i].text.replace(key, value)
    return document

def main():
    for name in os.listdir(old_file_path):
        print(name)
        old_file = old_file_path + name
        new_file = new_file_path + name
        if old_file.split(".")[1] == 'docx':
            document = Document(old_file)
            document = check_and_change(document, replace_dict)
            document.save(new_file)
        print("^"*30)

if __name__ == '__main__':
    main()
# ***********************************************************************************************************
# import os
# import comtypes.client
# def get_path():
#     # 指定路径
#     path = 'c:/Users/XUHAIJUN/SynologyDrive/English/friends/1/'
#     # 获取所有文件名的列表
#     filename_list = os.listdir(path)
#     # 获取所有word文件名列表
#     wordname_list = [filename for filename in filename_list \
#                      if filename.endswith((".doc", ".docx"))]
#     for wordname in wordname_list:
#         # 分离word文件名称和后缀，转化为pdf名称
#         pdfname = os.path.splitext(wordname)[0] + '.pdf'
#         # 如果当前word文件对应的pdf文件存在，则不转化
#         if pdfname in filename_list:
#             continue
#         # 拼接 路径和文件名
#         wordpath = os.path.join(path, wordname)
#         pdfpath = os.path.join(path, pdfname)
#         #生成器
#         yield wordpath,pdfpath
# def convert_word_to_pdf():
#     word = comtypes.client.CreateObject("Word.Application")
#     word.Visible = 0
#     for wordpath,pdfpath in get_path():
#         newpdf = word.Documents.Open(wordpath)
#         newpdf.SaveAs(pdfpath, FileFormat=17)
#         newpdf.Close()
# if __name__ == "__main__":
#     convert_word_to_pdf()