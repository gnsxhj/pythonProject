#!/usr/bin/env python
# -*- coding:utf-8 -*-
# coding=utf-8
from pdfminer.pdfparser import PDFParser, PDFDocument
from pdfminer.pdfinterp import PDFResourceManager, PDFPageInterpreter
from pdfminer.layout import LAParams
from pdfminer.converter import PDFPageAggregator
from docx import Document
import warnings
import os
import time
import glob
import fitz
import re
J = True
while J:
    print('1.PDF_转_Word\n2.PDF图片提取\n3.退出系统\n')
    print("""温馨提示:选择功能之前请务必将你所要处理的文件与该程序放在同一文件夹中!\n""")
    choice = int(input('请输入数字选择对应的功能:'))
    if choice == 1:
        print('小可爱，你好！欢迎使用PDF_转_Word程序！')
        time.sleep(1.5)
        print(
            '———————————————————————————————————Welcome to the program！——————————————————————————————————————————————')
        time.sleep(1.5)

        pdf_list = glob.glob('c:/Users/XUHAIJUN/Documents/SVW_Consulter_Strategie_2023年/人员简历/PDF/*.pdf') # 查看同文件夹下的pdf文件数
        print(u'共发现%s个pdf文件' % len(pdf_list))
        print(u'正在处理............')
        print(pdf_list)
        for l in iter(pdf_list):
            file_name = os.open(l, os.O_RDWR)
            document = Document()
            warnings.filterwarnings("ignore")

            def pdf2word():
                fn = open(file_name, 'rb')
                parser = PDFParser(fn)
                doc = PDFDocument()
                parser.set_document(doc)
                doc.set_parser(parser)
                resource = PDFResourceManager()
                laparams = LAParams()
                device = PDFPageAggregator(resource, laparams=laparams)
                interpreter = PDFPageInterpreter(resource, device)
                for i in doc.get_pages():
                    interpreter.process_page(i)
                    layout = device.get_result()
                    for out in layout:
                        if hasattr(out, "get_text"):
                            content = out.get_text().replace(u'\xa0', u' ')
                            document.add_paragraph(
                                content, style='ListBullet'
                            )
                    document.save(l + '.docx')
            pdf2word()
        print('处理完成')
        break
    elif choice == 2:
        print('小可爱，你好！欢迎使用PDF图片提取系统！')
        time.sleep(1.5)
        print(
            '———————————————————————————————————Welcome to the program！——————————————————————————————————————————————')
        time.sleep(1.5)

        def pdf2pic(path, pic_path):
            t0 = time.clock()  # 生成图片初始时间
            checkXO = r"/Type(?= */XObject)"  # 使用正则表达式来查找图片
            checkIM = r"/Subtype(?= */Image)"
            doc = fitz.open(path)  # 打开pdf文件
            imgcount = 0  # 图片计数
            lenXREF = doc._getXrefLength()  # 获取对象数量长度
            # 打印PDF的信息
            print("文件名:{}, 页数: {}, 对象: {}".format(path, len(doc), lenXREF - 1))
            # 遍历每一个对象
            for i in range(1, lenXREF):
                text = doc._getXrefString(i)  # 定义对象字符串
                isXObject = re.search(checkXO, text)  # 使用正则表达式查看是否是对象
                isImage = re.search(checkIM, text)  # 使用正则表达式查看是否是图片
                if not isXObject or not isImage:  # 如果不是对象也不是图片，则continue
                    continue
                imgcount += 1
                pix = fitz.Pixmap(doc, i)  # 生成图像对象
                new_name = "图片{}.png".format(imgcount)  # 生成图片的名称
                if pix.n < 5:  # 如果pix.n<5,可以直接存为PNG
                    pix.writePNG(os.path.join(pic_path, new_name))
                else:  # 否则先转换CMYK
                    pix0 = fitz.Pixmap(fitz.csRGB, pix)
                    pix0.writePNG(os.path.join(pic_path, new_name))
                    pix0 = None
                pix = None  # 释放资源
                t1 = time.clock()  # 图片完成时间
                print("运行时间:{}s".format(t1 - t0))
                print("提取了{}张图片".format(imgcount))
        if __name__ == '__main__':
            path = input('请输入需提取PDF文件路径:')
            pic_path = input('请输入提取图片保存路径:')
            # 创建保存图片的文件夹
            if os.path.exists(pic_path):
                print("文件夹已存在，不必重新创建！")
                pass
            else:
                os.mkdir(pic_path)
            pdf2pic(path, pic_path)
        break
    elif choice == 3:
        print('拜拜~欢迎下次光临！')
        break
    else:
        print('小伙子\小姐姐不要开车哦！我可不是傻子，请按照正确的流程输入！')
        print('——————————————————————❁—————————————————————————❁———————————————————————————❁————————————————————————')
        time.sleep(1.5)
        continue